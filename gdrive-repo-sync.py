#!/usr/bin/env python3
# /// script
# requires-python = ">=3.13"
# dependencies = [
#     "python-docx",
# ]
# ///
import os
import json
import subprocess
import fnmatch
import argparse
from pathlib import Path

from docx import Document

WORKSPACE_DIR = "workspace"
OUTPUT_DIR = "output"
GDRIVE_REMOTE = "gdrive"
MAX_FILE_LINES = 250


def run_cmd(cmd, cwd=None):
    return subprocess.run(cmd, cwd=cwd, capture_output=True, text=True, check=False).stdout.strip()


def get_gitignore_patterns(repo_path):
    gitignore_path = Path(repo_path) / ".gitignore"
    if not gitignore_path.exists():
        return []

    with open(gitignore_path, "r", encoding="utf-8", errors="ignore") as f:
        return [line.strip() for line in f if line.strip() and not line.startswith("#")]


def should_ignore_file(file_path, repo_path, gitignore_patterns, excluded_extensions):
    rel_path = os.path.relpath(file_path, repo_path)
    ext = os.path.splitext(file_path)[1].lower()

    if rel_path.startswith(".git/") or ext in excluded_extensions:
        return True

    for pattern in gitignore_patterns:
        if fnmatch.fnmatch(rel_path, pattern) or fnmatch.fnmatch(os.path.basename(rel_path), pattern):
            return True

    return False


def process_repo(repo_url, max_lines, excluded_extensions):
    repo_name = repo_url.split("/")[-1].replace(".git", "")
    repo_path = Path(WORKSPACE_DIR) / repo_name

    if not repo_path.exists():
        os.makedirs(WORKSPACE_DIR, exist_ok=True)
        run_cmd(["git", "clone", repo_url, str(repo_path)])
    else:
        run_cmd(["git", "pull", "-f"], cwd=str(repo_path))

    gitignore_patterns = get_gitignore_patterns(repo_path)
    doc = Document()

    doc.add_heading(f"Repository: {repo_name}", level=0)

    doc.add_heading("Table of Contents", level=1)
    toc = doc.add_paragraph()

    file_list = []

    for file_path in repo_path.glob('**/*'):
        if not file_path.is_file():
            continue

        if should_ignore_file(str(file_path), str(repo_path), gitignore_patterns, excluded_extensions):
            continue

        line_count = sum(1 for _ in open(file_path, "r", encoding="utf-8", errors="ignore"))
        if line_count > max_lines:
            continue

        rel_path = file_path.relative_to(repo_path)
        file_list.append(str(rel_path))

    for file_path in file_list:
        toc.add_run(f"• {file_path}\n")

    doc.add_page_break()

    for file_path_str in file_list:
        file_path = repo_path / file_path_str

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            file_content = f.read()

        doc.add_heading(file_path_str, level=1)

        p = doc.add_paragraph()
        p.add_run(file_content)

        doc.add_page_break()

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(str(Path(OUTPUT_DIR) / f"{repo_name}.docx"))

    return repo_name


def sync_to_gdrive():
    subprocess.run(
        [
            "rclone", "copy", "--drive-import-formats",
            "docx", OUTPUT_DIR, f"{GDRIVE_REMOTE}:repo-docs"
        ]
    )


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--repo-list", default="repos.json", help="Path to JSON file with repo URLs")
    parser.add_argument("--max-lines", type=int, default=MAX_FILE_LINES, help="Maximum number of lines per file")
    parser.add_argument(
        "--exclude", default=".svg,.ico,.png,.jpg,.jpeg,.gif,.bmp,.ttf,.woff,.woff2,.eot,.dll,.exe,.bin",
        help="Comma-separated list of file extensions to exclude"
    )
    parser.add_argument("--no-sync", action="store_true", help="Skip syncing to Google Drive")
    parser.add_argument("--debug", action="store_true", help="Enable debug output")
    args = parser.parse_args()

    excluded_extensions = set(
        ext.strip() if ext.startswith('.') else f'.{ext.strip()}' for ext in args.exclude.split(',')
    )

    os.makedirs(WORKSPACE_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    with open(args.repo_list, "r") as f:
        repos = json.load(f)

    for repo_url in repos:
        if args.debug:
            print(f"Processing {repo_url}")
        process_repo(repo_url, args.max_lines, excluded_extensions)

    if not args.no_sync:
        if args.debug:
            print("Syncing to Google Drive")
        sync_to_gdrive()


if __name__ == "__main__":
    main()